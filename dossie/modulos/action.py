# modulos/action.py
# -*- coding: utf-8 -*-
from __future__ import annotations
from io import BytesIO
from typing import Dict, Iterable, List, Optional, Tuple
import re

from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


# -------------------------
# Utilitários de parágrafos
# -------------------------
def _iter_paragraphs_in_cell(cell) -> Iterable[Paragraph]:
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                yield from _iter_paragraphs_in_cell(c)

def _iter_all_paragraphs(doc: _Document) -> Iterable[Paragraph]:
    # Corpo
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                yield from _iter_paragraphs_in_cell(c)
    # Header/Footer
    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                yield p
            for t in sec.header.tables:
                for row in t.rows:
                    for c in row.cells:
                        yield from _iter_paragraphs_in_cell(c)
        if sec.footer:
            for p in sec.footer.paragraphs:
                yield p
            for t in sec.footer.tables:
                for row in t.rows:
                    for c in row.cells:
                        yield from _iter_paragraphs_in_cell(c)

def _full_text_and_run_spans(paragraph: Paragraph) -> Tuple[str, List[Tuple[int, int]]]:
    full = []
    spans = []
    cur = 0
    for r in paragraph.runs:
        txt = r.text or ""
        full.append(txt)
        spans.append((cur, cur + len(txt)))
        cur += len(txt)
    return "".join(full), spans

def _replace_span_preserving_runs(paragraph: Paragraph, span: Tuple[int, int], replacement: str) -> None:
    start, end = span
    if start >= end or not paragraph.runs:
        return
    _, spans = _full_text_and_run_spans(paragraph)
    overlaps: List[Tuple[int, int, int]] = []
    for i, (rs, re_) in enumerate(spans):
        if rs < end and re_ > start:
            seg_start = max(start, rs)
            seg_end = min(end, re_)
            overlaps.append((i, seg_start - rs, seg_end - rs))

    rep = replacement or ""
    pos = 0
    for k, (i, rel_s, rel_e) in enumerate(overlaps):
        r = paragraph.runs[i]
        before = r.text[:rel_s]
        after = r.text[rel_e:]
        seg_len = rel_e - rel_s
        if k < len(overlaps) - 1:
            chunk = rep[pos: pos + seg_len]
            pos += len(chunk)
        else:
            chunk = rep[pos:]
        r.text = before + chunk + after


# -------------------------
# Regras de substituição
# -------------------------
_XS = r"(?P<slot>[xX]+)"
_RE_EMPRESA = re.compile(rf"empresa\s*:\s*{_XS}", re.IGNORECASE)
_RE_ANO = re.compile(rf"ano\s+de\s+(?:candidatura|condidatura)\s*:\s*{_XS}", re.IGNORECASE)
_RE_ESTADO_SUFFIX = re.compile(rf"{_XS}\s*-\s*brasil\b", re.IGNORECASE)

def _replace_slots_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full, _ = _full_text_and_run_spans(paragraph)
    low = full.lower()
    pending: List[Tuple[Tuple[int, int], str]] = []

    nome = (replacements.get("empresa") or "").strip()
    if nome:
        for m in _RE_EMPRESA.finditer(low):
            s, e = m.span("slot")
            pending.append(((s, e), nome))

    ano = (replacements.get("ano") or "").strip()
    if ano:
        for m in _RE_ANO.finditer(low):
            s, e = m.span("slot")
            pending.append(((s, e), ano))

    estado = (replacements.get("estado") or "").strip()
    if estado:
        for m in _RE_ESTADO_SUFFIX.finditer(low):
            s, e = m.span("slot")
            pending.append(((s, e), estado))

    for (s, e), rep in sorted(pending, key=lambda x: x[0][0], reverse=True):
        _replace_span_preserving_runs(paragraph, (s, e), rep)


# -------------------------
# Imagens — localizar/deletar
# -------------------------
def _local_name(tag: str) -> str:
    return tag.split('}', 1)[-1] if '}' in tag else tag

def _get_rel_attr(el, local: str) -> Optional[str]:
    return el.get(f"{{{R_NS}}}{local}") or el.get(local)

def _collect_images_with_containers_in_part(part, where: str) -> List[Dict]:
    items: List[Dict] = []
    root = part.element

    # DrawingML: remover o elemento <w:drawing>
    for draw in root.xpath('.//*[local-name()="drawing"]'):
        # tenta achar blip -> rid (opcional para ranking, não necessário para deletar)
        rid = None
        blips = draw.xpath('.//*[local-name()="blip"]')
        if blips:
            rid = _get_rel_attr(blips[0], "embed")
        items.append({
            "where": where,
            "container": draw,
            "rid": rid,
            "kind": "drawing",
        })

    # VML: remover o elemento <w:pict> (ou v:shape se não achar pict)
    for pict in root.xpath('.//*[local-name()="pict"]'):
        rid = None
        im = pict.xpath('.//*[local-name()="imagedata"]')
        if im:
            rid = _get_rel_attr(im[0], "id")
        items.append({
            "where": where,
            "container": pict,
            "rid": rid,
            "kind": "pict",
        })

    # Fallback direto: se houver <a:blip> solto sem <w:drawing>, tenta subir até um contêiner removível
    for blip in root.xpath('.//*[local-name()="blip"]'):
        container = None
        node = blip
        for _ in range(12):
            node = node.getparent()
            if node is None:
                break
            if _local_name(node.tag) in ("drawing", "pict", "r"):  # w:drawing, w:pict ou o run
                container = node
                break
        if container is not None:
            rid = _get_rel_attr(blip, "embed")
            items.append({
                "where": where,
                "container": container,
                "rid": rid,
                "kind": "blip",
            })

    # VML direto: <v:imagedata> sem pict
    for im in root.xpath('.//*[local-name()="imagedata"]'):
        container = None
        node = im
        for _ in range(12):
            node = node.getparent()
            if node is None:
                break
            if _local_name(node.tag) in ("pict", "shape", "r"):
                container = node
                break
        if container is not None:
            rid = _get_rel_attr(im, "id")
            items.append({
                "where": where,
                "container": container,
                "rid": rid,
                "kind": "imagedata",
            })

    return items

def _collect_all_images_with_containers(doc: _Document) -> List[Dict]:
    items: List[Dict] = []
    # Primeiro o BODY (capa costuma estar no corpo)
    items += _collect_images_with_containers_in_part(doc.part, "body")
    # Depois header e footer
    for sec in doc.sections:
        if sec.header:
            items += _collect_images_with_containers_in_part(sec.header.part, "header")
        if sec.footer:
            items += _collect_images_with_containers_in_part(sec.footer.part, "footer")
    return items

def _delete_element(el) -> bool:
    if el is None:
        return False
    parent = el.getparent()
    if parent is None:
        return False
    try:
        parent.remove(el)
        return True
    except Exception:
        return False


# -------------------------
# API pública
# -------------------------
def remover_logo_capa(docx_bytes: bytes) -> bytes:
    if not docx_bytes:
        raise ValueError("Documento .docx não fornecido ou vazio.")

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return docx_bytes

    items = _collect_all_images_with_containers(doc)
    if items:
        # Prioridade: body -> header -> footer; e, dentro de cada, o primeiro
        def score(it):
            s = 0
            where = it.get("where")
            if where == "body":
                s += 300
            elif where == "header":
                s += 200
            elif where == "footer":
                s += 100
            # tipos mais altos primeiro
            kind = it.get("kind")
            if kind == "drawing":
                s += 3
            elif kind == "pict":
                s += 2
            else:
                s += 1
            return s
        items.sort(key=score, reverse=True)
        _delete_element(items[0]["container"])

    out = BytesIO()
    try:
        doc.save(out)
        data = out.getvalue()
        return data if data else docx_bytes
    except Exception:
        return docx_bytes


def preencher_modelo(
    docx_bytes: bytes,
    nome_empresa: Optional[str],
    estado_extenso: Optional[str],
    ano_candidatura: Optional[str],
    logo_bytes: Optional[bytes] = None,  # ignorado por enquanto
) -> bytes:
    if not docx_bytes:
        raise ValueError("Documento .docx não fornecido ou vazio.")

    # 1) Substituições de texto
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return docx_bytes

    reps = {
        "empresa": (nome_empresa or "").strip(),
        "estado": (estado_extenso or "").strip(),
        "ano": (ano_candidatura or "").strip(),
    }
    for p in _iter_all_paragraphs(doc):
        _replace_slots_in_paragraph(p, reps)

    # 2) Deletar a primeira imagem (capa)
    items = _collect_all_images_with_containers(doc)
    if items:
        def score(it):
            s = 0
            where = it.get("where")
            if where == "body":
                s += 300
            elif where == "header":
                s += 200
            elif where == "footer":
                s += 100
            kind = it.get("kind")
            if kind == "drawing":
                s += 3
            elif kind == "pict":
                s += 2
            else:
                s += 1
            return s
        items.sort(key=score, reverse=True)
        _delete_element(items[0]["container"])

    # 3) Salvar com fallback seguro
    out = BytesIO()
    try:
        doc.save(out)
        data = out.getvalue()
        return data if data else docx_bytes
    except Exception:
        return docx_bytes