# app.py
from __future__ import annotations
import re, unicodedata
from datetime import date
from io import BytesIO
import streamlit as st
import pandas as pd

# ========================== Config ==========================
st.set_page_config(page_title="Preencher Capa (PDF → DOCX)", page_icon="🪩", layout="centered")
st.title("🪩 Preencher Capa do Recurso (PDF → DOCX)")
st.caption("1) Envie PDFs (Parecer, RFG/CNPJ, Form PD) e o logo • 2) Ajuste dados • 3) Envie o modelo .docx e gere")

# ====================== Util: garantir bytes ======================
def _ensure_bytes(x, label: str = "resultado") -> bytes:
    if isinstance(x, bytes): return x
    if isinstance(x, bytearray): return bytes(x)
    try: return bytes(x)
    except Exception: raise ValueError(f"{label} não é binário (tipo: {type(x)})")

# ======================= Cache helpers (lazy imports) =======================
@st.cache_data(show_spinner=False)
def _pdf_text(raw: bytes) -> str:
    from modulos import extrator as XP
    return XP.extract_text_from_pdf(raw)

@st.cache_data(show_spinner=False)
def _parecer_fields(text: str) -> dict:
    from modulos import extrator as XP
    return XP.extract_fields(text)

@st.cache_data(show_spinner=False)
def _rfg_fields(text: str) -> dict:
    from modulos import extrator_rfg as XR
    return XR.extract_rfg_fields(text)

@st.cache_data(show_spinner=False)
def _form_fields(raw: bytes, item_index: int | None = None) -> dict:
    from modulos.extrator_form import extract_form_fields
    return extract_form_fields(raw, item_index=item_index)

@st.cache_data(show_spinner=False)
def _form_md(fields: dict) -> str:
    from modulos.extrator_form import as_markdown
    return as_markdown(fields)

@st.cache_data(show_spinner=False)
def _form_count(raw: bytes) -> int:
    from modulos.extrator_form import count_items
    return count_items(raw)

@st.cache_data(show_spinner=False)
def _try_formpd_any(raw: bytes) -> dict | None:
    try:
        from modulos.extrator_form import extract_form_fields
        f = extract_form_fields(raw)
    except Exception:
        return None
    score = sum(1 for k in (
        "pb_pa_ou_de","area_do_projeto","natureza",
        "destaque_elemento_novo","barreira_desafio_tecnologico","metodologia_metodos"
    ) if f.get(k))
    return f if score >= 2 else None

# === RH ===
@st.cache_data(show_spinner=False)
def _rh_extract(raw: bytes) -> dict:
    from modulos.extrator_rh import extract_rh_entries
    res = extract_rh_entries(raw)
    return {"items": res.items, "total_text": res.total_text, "section_bounds": res.section_bounds}

# ===================== Detector por score + validação RFG =====================
ANCH = re.compile(r"\b3\.\s*(?:1\.\s*(?:1|2)\s*)?\b", re.IGNORECASE)
RFG_TOKENS = tuple(s.upper() for s in [
    "Cadastro Nacional da Pessoa Jurídica","Comprovante de Inscrição","Situação Cadastral",
    "Nome Empresarial","Natureza Jurídica","Atividade Econômica Principal","Data de Abertura"
])
FORM_LABELS = ("PB, PA OU DE","ÁREA DO PROJETO","AREA DO PROJETO","NATUREZA")

def _score_parecer(t: str) -> int:
    return (3 if "PARECER TÉCNICO" in t else 0) + (1 if re.search(r"\bPARECER\s+T[ÉE]CNICO\b", t) else 0)

def _score_rfg(t: str) -> int:
    return sum(1 for tok in RFG_TOKENS if tok in t)

def _score_form(t: str) -> int:
    return len(ANCH.findall(t)) + sum(1 for lb in FORM_LABELS if lb in t)

def detect_tipo(text: str) -> str:
    t = text.upper(); cand: list[tuple[str,int]] = []
    sp, sr, sf = _score_parecer(t), _score_rfg(t), _score_form(t)
    if sp >= 2: cand.append(("Parecer", sp))
    if sr >= 3: cand.append(("RFG", sr))
    if sf >= 3 or len(ANCH.findall(t)) >= 2: cand.append(("FormPD", sf))
    if not cand: return "Desconhecido"
    cand.sort(key=lambda x: x[1], reverse=True)
    return cand[0][0]

UF_SET = set("AC AL AP AM BA CE DF ES GO MA MT MS MG PA PB PR PE PI RJ RN RS RO RR SC SP SE TO".split())
CEP_RX = re.compile(r"^\d{5}\-?\d{3}$")

def _valid_rfg(r: dict) -> bool:
    filled = sum(1 for k in ("Logradouro","Número","CEP","Município","UF") if (r.get(k) or "").strip())
    return filled >= 3 and (r.get("UF","" ).strip().upper() in UF_SET) and bool(CEP_RX.match((r.get("CEP","" ).strip())))

# ============================ Estado inicial ============================
ss = st.session_state
ss.setdefault("parecer_data", {})
ss.setdefault("rfg_data", {})
ss.setdefault("form_fields", None)
ss.setdefault("has_inputs", False)
ss.setdefault("formpd_pdf", None)
ss.setdefault("logo_bytes", None)
ss.setdefault("rh_items", [])
ss.setdefault("rh_total_line", None)
ss.setdefault("item_escolhido", 1)
ss.setdefault("rh_xlsx_bytes", None)

# ============================= 1) Upload PDFs + Logo =============================
st.subheader("1) Envie os PDFs (Parecer, RFG/CNPJ e Form PD) e o Logo")

c_pdf, c_logo = st.columns([3,2])
with c_pdf:
    pdf_files = st.file_uploader("Selecione os PDFs (pode enviar todos de uma vez)", type=["pdf"], accept_multiple_files=True, key="pdfs_input")
with c_logo:
    logo_file = st.file_uploader("Logo da empresa (PNG/JPG)", type=["png","jpg","jpeg","bmp","gif"], key="logo_input")
    if logo_file is not None:
        ss.logo_bytes = logo_file.read(); st.caption("Logo carregado na sessão.")
    elif ss.logo_bytes:
        st.caption("Logo mantido da sessão.")

c1, c2 = st.columns(2)
with c1:
    processar = st.button("📥 Processar PDFs", key="processar_pdfs_btn")
with c2:
    if st.button("🔄 Limpar sessão", key="clean_session_btn"):
        st.cache_data.clear()
        for k in ("parecer_data","rfg_data","form_fields","has_inputs","formpd_pdf","logo_bytes","rh_items","rh_total_line","item_escolhido","rh_xlsx_bytes"):
            ss.pop(k, None)
        st.experimental_rerun()

if pdf_files and processar:
    parecer, rfg, fpd = {}, {}, None
    formpd_pdf_local = None
    detected: list[tuple[str,str]] = []
    with st.spinner("Lendo e classificando PDFs..."):
        for f in pdf_files:
            raw = f.read()
            if not raw:
                detected.append((f.name, "Arquivo vazio")); continue
            text = _pdf_text(raw)
            tipo = detect_tipo(text); tipo_final = tipo

            if tipo == "Parecer" and not parecer:
                fields = _parecer_fields(text)
                parecer = {
                    "empresa": fields.get("Nome da empresa",""),
                    "cnpj": fields.get("CNPJ",""),
                    "parecer_contestacao": fields.get("Nº do Parecer Técnico da Contestação",""),
                    "parecer_tecnico": fields.get("Nº do Parecer Técnico",""),
                    "ano_base": fields.get("Ano-Base"),
                }
            elif tipo == "RFG" and not rfg:
                r = _rfg_fields(text)
                if _valid_rfg(r):
                    rfg = {"logradouro": r.get("Logradouro",""), "numero": r.get("Número",""), "cep": r.get("CEP",""), "municipio": r.get("Município",""), "uf": r.get("UF","")}
                else:
                    tr = _try_formpd_any(raw)
                    if tr and fpd is None:
                        fpd, tipo_final = tr, "FormPD (fallback via RFG inválido)"; formpd_pdf_local = raw
                    else:
                        tipo_final = "Desconhecido (RFG inválido)"
            elif tipo == "FormPD" and fpd is None:
                try:
                    fpd = _form_fields(raw, item_index=None)
                    formpd_pdf_local = raw
                except Exception as e:
                    st.error(f"Form PD: {e}")

            if tipo == "Desconhecido" and fpd is None:
                tr = _try_formpd_any(raw)
                if tr:
                    fpd, tipo_final = tr, "FormPD (fallback)"; formpd_pdf_local = raw

            try:
                if (tipo == "FormPD" or str(tipo_final).startswith("FormPD")) and not ss.get("rh_items"):
                    rh = _rh_extract(raw)
                    if rh and rh.get("items"):
                        ss.rh_items = rh["items"]; ss.rh_total_line = rh.get("total_text")
            except Exception:
                pass

            detected.append((f.name, tipo_final))

        with st.expander("Arquivos identificados"):
            for n,t in detected:
                st.write(f"• **{n}** → {t}")
            if ss.get("rh_items"):
                st.write(f"• Relação de RH: {len(ss.rh_items)} item(ns) identificado(s).")

        ss.parecer_data = parecer or ss.parecer_data
        ss.rfg_data = rfg or ss.rfg_data
        ss.form_fields = fpd if fpd is not None else ss.form_fields
        ss.formpd_pdf = formpd_pdf_local or ss.formpd_pdf
        if ss.form_fields is not None: ss.form_fields["numero_projeto"] = "1"

        if ss.get("formpd_pdf") and not ss.get("rh_items"):
            try:
                rh = _rh_extract(ss.formpd_pdf)
                if rh and rh.get("items"):
                    ss.rh_items = rh["items"]; ss.rh_total_line = rh.get("total_text")
            except Exception:
                pass

        ss.has_inputs = any([ss.parecer_data, ss.rfg_data, ss.form_fields])

if not ss.has_inputs:
    st.info("Envie e processe os PDFs para continuar."); st.stop()

# ===================== Configuração dinâmica: Item do Form PD =====================
st.subheader("⚙️ Configuração da extração do Form PD (Item)")
if ss.get("formpd_pdf"):
    try: n_items = _form_count(ss.formpd_pdf)
    except Exception: n_items = 1
    st.caption(f"Detectados **{n_items}** item(ns) no Form PD.")

    def _make_exclusive_cb(idx: int, total: int):
        def _cb():
            for j in range(1, total+1):
                if j != idx: st.session_state[f"form_item_{j}"] = False
            st.session_state["item_escolhido"] = idx
        return _cb

    cols = st.columns(n_items)
    if "item_escolhido" not in ss or not (1 <= ss.item_escolhido <= n_items): ss.item_escolhido = 1
    for i in range(1, n_items+1):
        default = (i == ss.item_escolhido)
        with cols[i-1]:
            st.checkbox(f"Item {i}", key=f"form_item_{i}", value=default if f"form_item_{i}" not in st.session_state else st.session_state[f"form_item_{i}"], on_change=_make_exclusive_cb(i, n_items))

    apply_col, _ = st.columns([1,3])
    with apply_col:
        if st.button("Aplicar Item selecionado"):
            sel = ss.get("item_escolhido", 1)
            new_fields = _form_fields(ss.formpd_pdf, item_index=sel if n_items > 1 else None)
            if new_fields:
                new_fields["numero_projeto"] = "1"; ss.form_fields = new_fields
                st.success(f"Campos re-extraídos do **Item {sel}**.")
            else:
                st.warning("Não foi possível re-extrair os campos para o Item selecionado.")
else:
    st.info("Após processar os PDFs, selecione aqui o Item (1..N) do Form PD a extrair.")

st.divider()

# =========================== 2) Revisão/Edição ===========================
st.subheader("2) Revisar/editar variáveis")
pd_data = ss.parecer_data or {}; rd = ss.rfg_data or {}
col1, col2 = st.columns(2)
with col1:
    empresa = st.text_input("Empresa (Razão Social)", pd_data.get("empresa",""))
    cnpj = st.text_input("CNPJ", pd_data.get("cnpj",""))
    parecer_ct = st.text_input("Parecer Técnico da Contestação nº", pd_data.get("parecer_contestacao",""))
    parecer_tec = st.text_input("Parecer Técnico nº", pd_data.get("parecer_tecnico",""))
    ano_base = st.number_input("Ano-base", 2005, 2100, int(pd_data.get("ano_base") or date.today().year), 1)
    lote = st.text_input("Número do lote (opcional)", "")
with col2:
    logradouro = st.text_input("Logradouro", rd.get("logradouro",""))
    numero = st.text_input("Número (endereço)", rd.get("numero",""))
    cep = st.text_input("CEP", rd.get("cep",""))
    municipio = st.text_input("Município", rd.get("municipio",""))
    uf = st.text_input("UF (ex.: GO)", rd.get("uf",""))

from modulos.action import UF_TO_ESTADO
st.caption("Local/Data e assinatura")
c3, c4, c5, c6 = st.columns(4)
with c3: estado_nome = st.text_input("Estado (extenso)", UF_TO_ESTADO.get((uf or "").strip().upper(), ""))
with c4: dia = st.number_input("Dia", 1, 31, date.today().day, 1)
with c5: mes = st.number_input("Mês", 1, 12, date.today().month, 1)
with c6: ano_data = st.number_input("Ano (data)", 2005, 2100, date.today().year, 1)
responsavel = st.text_input("Responsável (FormPD)", "")

if parecer_ct and parecer_tec and parecer_ct.strip()==parecer_tec.strip():
    st.warning("Os números do **Parecer da Contestação** e do **Parecer Técnico** estão iguais.")

st.divider()

# =========================== Trechos do Form PD ===========================
if ss.form_fields:
    ss.form_fields["numero_projeto"] = "1"
    if ss.formpd_pdf and (("nome_projeto" not in ss.form_fields) or ("numero_projeto" not in ss.form_fields)):
        try:
            from modulos.extrator_form import extract_form_fields as _ef
            fresh = _ef(ss.formpd_pdf, item_index=ss.get("item_escolhido", None))
            if fresh.get("nome_projeto") is not None:
                ss.form_fields["nome_projeto"] = fresh.get("nome_projeto"); ss.form_fields["numero_projeto"] = "1"
        except Exception:
            pass

    st.subheader("Trechos do Form PD")
    md = _form_md(ss.form_fields)
    prepend = []
    np_val = ss.form_fields.get("nome_projeto"); nr_val = ss.form_fields.get("numero_projeto") or "1"
    if np_val and "**Nome do Projeto / Atividade (3.1.1)**" not in md:
        prepend.append(f"**Nome do Projeto / Atividade (3.1.1):**\n\n{np_val}\n")
    if "**Número do Projeto (3.1.2)**" not in md:
        prepend.append(f"**Número do Projeto (3.1.2):**\n\n{nr_val}\n")
    if prepend: md = "\n".join(prepend) + "\n" + md
    st.markdown(md)
    st.download_button("⬇️ Baixar trechos (.md)", data=md.encode("utf-8"), file_name="form_pd_trechos.md", mime="text/markdown")
else:
    st.info("Form PD não detectado. Inclua o Form PD no passo 1 para extrair os trechos.")

# =========================== Planilha RH (Relação de Recursos Humanos) ===========================
st.divider(); st.subheader("Planilha — Relação de Recursos Humanos")

def _fmt_valor(v):
    if isinstance(v, float):
        s = f"{v:,.2f}"; s = s.replace(",","X").replace(".",",").replace("X",".")
        return f"R$ {s}"
    return ""

rows = []; total_val = 0.0; items = ss.get("rh_items") or []
for it in items:
    valor = it.get("valor_rs")
    if isinstance(valor, float): total_val += valor
    rows.append({
        "Item": it.get("item_numero",""),
        "CPF": it.get("cpf_formatado") or it.get("cpf") or "",
        "Nome": it.get("nome") or "",
        "Titulação": it.get("titulacao") or "",
        "Total horas (Anual)": it.get("total_horas_anual") or "",
        "Dedicação": it.get("dedicacao") or "",
        "Valor (R$)": _fmt_valor(valor),
    })

try:
    cols = ["Item","CPF","Nome","Titulação","Total horas (Anual)","Dedicação","Valor (R$)"]
    df = pd.DataFrame(rows, columns=cols)
    st.dataframe(df, use_container_width=True, hide_index=True)
except Exception:
    if rows: st.dataframe(rows, use_container_width=True, hide_index=True)
    else:
        st.table({"Item":[],"CPF":[],"Nome":[],"Titulação":[],"Total horas (Anual)":[],"Dedicação":[],"Valor (R$)":[]})

# Download planilha modelo
if items:
    try:
        from types import SimpleNamespace
        import tempfile, os
        from modulos.extrator_rh import export_rh_to_model_sheet
        _res_like = SimpleNamespace(items=items, total_text=ss.get("rh_total_line"), section_bounds=(-1,-1), model="")
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tf:
            tmp_path = tf.name
        try:
            export_rh_to_model_sheet(_res_like, tmp_path, sheet_name="Planilha1", write_total_text_cell=None)
            with open(tmp_path, "rb") as f: xbytes = f.read()
        finally:
            try: os.remove(tmp_path)
            except: pass
        st.download_button("⬇️ Baixar planilha modelo (XLSX)", data=xbytes, file_name="Contestacao_RH.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_rh_xlsx")
    except Exception:
        st.warning("Não foi possível gerar a planilha XLSX no modelo. Verifique se 'openpyxl' está instalado.")

# Uploader da planilha reimportada (com Justificativa)
up_col1, _ = st.columns([1,3])
with up_col1:
    rh_xlsx_up = st.file_uploader("Reimportar planilha RH (XLSX)", type=["xlsx"], key="rh_xlsx_reimport", help="Envie a mesma planilha baixada, com a coluna 'Justificativa' preenchida.")
if rh_xlsx_up is not None:
    try:
        ss.rh_xlsx_bytes = rh_xlsx_up.getvalue()
        df_up = pd.read_excel(rh_xlsx_up, engine="openpyxl")
        lower_map = {str(c).lower().strip(): c for c in df_up.columns}
        just_col = lower_map.get("justificativa")
        ss.setdefault("rh_just", {})
        if not just_col:
            st.warning("A planilha enviada não contém a coluna 'Justificativa'. Verifique e envie novamente.")
        else:
            key_item = lower_map.get("item"); key_cpf = lower_map.get("cpf"); just_map = {}
            if key_item and key_cpf:
                for _, row in df_up.iterrows():
                    key = (str(row.get(key_item, "")).strip(), str(row.get(key_cpf, "")).strip())
                    just_map[key] = str(row.get(just_col, "") or "").strip()
                ss.rh_just = {"type":"tuple_key","keys":("Item","CPF"),"map":just_map}
                st.success(f"Justificativas importadas por chave (Item, CPF): {sum(1 for v in just_map.values() if v)} registros com texto.")
            else:
                for idx, row in df_up.iterrows():
                    just_map[int(idx)] = str(row.get(just_col, "") or "").strip()
                ss.rh_just = {"type":"row_index","map":just_map}
                st.success(f"Justificativas importadas por índice de linha: {sum(1 for v in just_map.values() if v)} registros com texto.")
            non_empty = [(k,v) for k,v in ss.rh_just["map"].items() if v][:5]
            if non_empty:
                st.caption("Exemplos de justificativas lidas:")
                for k,v in non_empty:
                    st.write(f"• {k}: {v[:120]}{'...' if len(v) > 120 else ''}")
    except Exception as e:
        st.error("Falha ao ler a planilha enviada."); st.exception(e)

colA, colB = st.columns([2,1])
with colA:
    st.caption(f"TOTAL (capturado do formulário): {ss.rh_total_line or '—'}")
with colB:
    st.caption(f"Soma calculada: {_fmt_valor(total_val)}")
if not items:
    st.info("RH não extraído")

# ========================= 3) DOCX: gerar e baixar =========================
st.subheader("3) Envie o modelo Word (.docx) e gere")
docx_file = st.file_uploader("Modelo Word (DOCX)", type=["docx"], key="docx_model_input")

from docx import Document
from docx.text.paragraph import Paragraph

def _norm_lower(s: str) -> str:
    s = s or ""; return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)).lower()

def _all_paragraphs(doc: Document) -> list[Paragraph]:
    pars: list[Paragraph] = []
    pars.extend(doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                pars.extend(cell.paragraphs)
    for section in doc.sections:
        if section.header:
            pars.extend(section.header.paragraphs)
            for t in section.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        pars.extend(c.paragraphs)
        if section.footer:
            pars.extend(section.footer.paragraphs)
            for t in section.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        pars.extend(c.paragraphs)
    return pars

def _quick_hit_count(doc_bytes: bytes, fields: dict) -> int:
    if not fields: return 0
    try:
        doc = Document(BytesIO(doc_bytes))
        text = "\n".join(p.text for p in _all_paragraphs(doc)).casefold()
    except Exception:
        return 0
    samples = []
    for k in ("pb_pa_ou_de","natureza","atividade_continua","data_inicio_atividade","previsao_termino","nome_projeto","nome_atividade_pdi","numero_projeto"):
        v = (fields.get(k) or "").strip()
        if not v: continue
        s = v.replace("\n"," ").strip(); samples.append((s[:50] if len(s)>=20 else s[:40]).casefold())
    return sum(1 for s in samples if s and s in text)

FIELD_RX = [
    ("destaque_elemento_novo", re.compile(r"destaque\s+o\s+elemento.*?novo", re.IGNORECASE | re.DOTALL), "next"),
    ("barreira_desafio_tecnologico", re.compile(r"(?:barreira|desafio).{0,40}?tecnol", re.IGNORECASE | re.DOTALL), "next"),
    ("metodologia_metodos", re.compile(r"metodologia\s*/?\s*m[eé]todos", re.IGNORECASE), "next"),
    ("pb_pa_ou_de", re.compile(r"\bpb,\s*pa\s*e\s*de\b", re.IGNORECASE), "inline"),
    ("natureza", re.compile(r"natureza\s+do\s+projeto.*formpd", re.IGNORECASE | re.DOTALL), "inline"),
    ("atividade_continua", re.compile(r"atividade\s*é?\s*cont[ií]nua", re.IGNORECASE), "inline"),
    ("data_inicio_atividade", re.compile(r"data\s+de\s+in[ií]cio\s+da\s+atividade", re.IGNORECASE), "next"),
    ("previsao_termino", re.compile(r"previs[aã]o\s+de\s+t[ée]rmino", re.IGNORECASE), "next"),
]

from docx.text.paragraph import Paragraph as _Paragraph

def _replace_inline_after_colon(p: _Paragraph, value: str) -> bool:
    runs = p.runs
    if not runs:
        p.add_run(": " + value); return True
    col_i, col_pos = -1, -1
    for i, r in enumerate(runs):
        t = r.text or ""; pos = t.find(":")
        if pos >= 0: col_i, col_pos = i, pos; break
    if col_i == -1:
        runs[-1].text = (runs[-1].text or "").rstrip() + ": "; p.add_run(value); return True
    runs[col_i].text = (runs[col_i].text or "")[:col_pos+1]
    if col_i+1 < len(runs):
        runs[col_i+1].text = " " + value
        for k in range(col_i+2, len(runs)): runs[k].text = ""
    else:
        p.add_run(" " + value)
    return True

def _replace_next_paragraph(p_label: _Paragraph, following: list[_Paragraph], value: str) -> bool:
    for p_next in following:
        if not p_next.text.strip() or p_next.text.strip().startswith("`"):
            if not p_next.runs: p_next.add_run("")
            for i in range(1, len(p_next.runs)): p_next.runs[i].text = ""
            p_next.runs[0].text = value; return True
        else:
            if not p_next.runs: p_next.add_run("")
            for i in range(1, len(p_next.runs)): p_next.runs[i].text = ""
            p_next.runs[0].text = value; return True
    return _replace_inline_after_colon(p_label, value)

def _apply_formpd_section_compat(docx_bytes: bytes, fields: dict | None) -> tuple[bytes, int, list[tuple[str,bool]]]:
    if not fields: return (docx_bytes, 0, [])
    try:
        doc = Document(BytesIO(docx_bytes)); pars = _all_paragraphs(doc)
        applied = 0; audit: list[tuple[str,bool]] = []
        for key, pat, mode in FIELD_RX:
            val = (fields.get(key) or "").strip()
            if not val: audit.append((key, False)); continue
            found_idx = -1
            for i, p in enumerate(pars):
                if pat.search(_norm_lower(p.text)): found_idx = i; break
            if found_idx == -1: audit.append((key, False)); continue
            ok = _replace_inline_after_colon(pars[found_idx], val) if mode=="inline" else _replace_next_paragraph(pars[found_idx], pars[found_idx+1:found_idx+4], val)
            audit.append((key, ok)); applied += int(ok)
            if key == "previsao_termino": break
        out = BytesIO(); doc.save(out); out.seek(0)
        final = out.read(); return (final if isinstance(final, (bytes,bytearray)) else docx_bytes, applied, audit)
    except Exception:
        return (docx_bytes, 0, [])

# ------- Gerar DOCX -------
if docx_file:
    ok_basic = all((empresa or '').strip() for empresa in (empresa, cnpj, municipio, uf))
    if not ok_basic:
        st.error("Preencha **Empresa, CNPJ, Município e UF** antes de gerar.")
    elif st.button("⚙️ Gerar DOCX", key="gerar_docx_btn"):
        try:
            from modulos.action import build_capa_context, fill_first_page
            from modulos.action_logo import insert_logo
            from modulos.action_form import fill_formpd_section
            from modulos.action_rh import locate_relacao_rh, fill_exclusives_from_sheet

            base = _ensure_bytes(docx_file.read(), "DOCX")

            try:
                found_rh, meta_rh = locate_relacao_rh(base)
                if found_rh:
                    if meta_rh.get("exclusivos_found"): st.success("Relação de RH localizada • Recursos Humanos Exclusivos: localizado")
                    else: st.info("Relação de RH localizada • Recursos Humanos Exclusivos: não encontrado")
                else:
                    st.info("Relação de RH não encontrada")
            except Exception:
                pass

            ctx = build_capa_context(
                empresa=empresa, cnpj=cnpj, parecer_contestacao=parecer_ct, parecer_tecnico=parecer_tec,
                ano_base=int(ano_base), logradouro=logradouro, numero=numero, cep=cep, municipio=municipio,
                uf=uf, lote=lote, estado_nome=estado_nome, dia=int(dia), mes=int(mes), ano_data=int(ano_data),
                responsavel=responsavel
            )

            tmp = _ensure_bytes(fill_first_page(base, ctx), "fill_first_page")
            if ss.get("logo_bytes"):
                tmp = insert_logo(tmp, ss.logo_bytes, max_width_cm=10.0, prefer_header=True, center=True)

            try:
                if ss.get("rh_xlsx_bytes"):
                    tmp, _meta_fill = fill_exclusives_from_sheet(tmp, ss.rh_xlsx_bytes, sheet_name="Planilha1")
                    st.info("Recursos Humanos Exclusivos: preenchidos/replicados a partir da planilha.")
            except Exception as _e:
                st.warning(f"RH Exclusivo: falha ao preencher a partir da planilha: {_e}")

            out = fill_formpd_section(tmp, ss.form_fields) if ss.form_fields else tmp
            final_bytes = _ensure_bytes(out if out else tmp, "arquivo final")

            hits_before = _quick_hit_count(tmp, ss.form_fields or {})
            hits_after = _quick_hit_count(final_bytes, ss.form_fields or {})
            applied_official = (hits_after > hits_before)

            audit = []
            if ss.form_fields and not applied_official:
                final_bytes, applied_alt, audit = _apply_formpd_section_compat(final_bytes, ss.form_fields)
                hits_after = _quick_hit_count(final_bytes, ss.form_fields or {})
                if applied_alt == 0 and hits_after <= hits_before:
                    st.warning("⚠️ O template não casou com os rótulos esperados. Gere o arquivo (capa ok) e me envie um print para ajustar as regex.")
                else:
                    st.info("✅ Textos do Form PD aplicados com compatibilidade de template.")

            st.download_button("⬇️ Baixar DOCX preenchido", data=final_bytes, file_name="recurso_capa_prog_pd&i_preenchido.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            with st.expander("Diagnóstico de aplicação (debug)"):
                st.write(f"Hits antes: {hits_before} \nHits depois: {hits_after} \nOficial aplicou? {applied_official}")
                if audit:
                    st.write("Campos (compatibilidade):")
                    for k, okk in audit: st.write(f"• {k}: {'ok' if okk else 'não encontrado'}")
                st.json(ctx)
        except Exception as e:
            st.error("Falha ao gerar o DOCX."); st.exception(e)
else:
    st.info("Envie o modelo Word (.docx) para continuar.")
