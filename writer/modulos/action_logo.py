# modulos/action_logo.py
from __future__ import annotations
from io import BytesIO
from typing import Iterable, Optional
import unicodedata
import re

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------------------------
# Normalização & util
# -------------------------------------------
def _norm_ascii_lower(s: str) -> str:
    s = (s or "").replace("\u00A0", " ")           # NBSP -> espaço comum
    s = s.replace("`", "").replace("´", "").strip()
    return "".join(c for c in unicodedata.normalize("NFKD", s)
                   if not unicodedata.combining(c)).lower()

def _all_body_paragraphs(doc: Document) -> list[Paragraph]:
    pars = list(doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                pars.extend(cell.paragraphs)
    return pars

def _all_header_paragraphs(doc: Document) -> list[Paragraph]:
    pars: list[Paragraph] = []
    if not doc.sections:
        return pars
    hdr = doc.sections[0].header
    if hdr:
        pars.extend(list(hdr.paragraphs))
        for tb in hdr.tables:
            for row in tb.rows:
                for cell in row.cells:
                    pars.extend(cell.paragraphs)
    return pars

def _available_width_cm(doc: Document, pad_cm: float = 0.0) -> float:
    """
    Largura útil (página - margens) em cm na 1ª seção, menos um 'padding' opcional.
    """
    sec = doc.sections[0]
    emu_avail = sec.page_width - (sec.left_margin + sec.right_margin)
    # 1 cm = 360000 EMUs
    avail_cm = float(emu_avail) / 360000.0
    return max(1.0, avail_cm - pad_cm)

def _clear_runs(p: Paragraph):
    if not p.runs:
        p.add_run("")
        return
    for r in p.runs:
        r.text = ""

def _force_following_runs_black(p: Paragraph, from_index: int = 0):
    for i, r in enumerate(p.runs):
        if i < from_index:
            continue
        try:
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except Exception:
            pass

# -------------------------------------------
# Busca por placeholder
# -------------------------------------------
PLACEHOLDER_TOKENS = (
    "logo da empresa",
    "logo da companhia",
    "logomarca da empresa",
)

# Limites para encerrarmos a varredura na 1ª página (heurística segura)
BODY_STOP_TOKENS = (
    "recurso administrativo",
    "sumário", "sumario",
    "programa e atividade de pd&i",
)

def _find_placeholder_in_pars(pars: Iterable[Paragraph]) -> Optional[Paragraph]:
    for p in pars:
        t = _norm_ascii_lower(p.text)
        if any(tok == t for tok in PLACEHOLDER_TOKENS):
            return p
        # também aceita se o texto do parágrafo for *apenas* o placeholder entre crases ou com espaçamento
        t_compact = re.sub(r"\s+", " ", t).strip(" :;-–")
        if any(tok == t_compact for tok in PLACEHOLDER_TOKENS):
            return p
    return None

def _find_placeholder_in_body_first_page(doc: Document) -> Optional[Paragraph]:
    pars = _all_body_paragraphs(doc)
    for p in pars:
        t = _norm_ascii_lower(p.text)
        if any(tok == t or tok == re.sub(r"\s+", " ", t).strip(" :;-–") for tok in PLACEHOLDER_TOKENS):
            return p
        # corta varredura quando chegamos em um título forte da página seguinte
        if any(stop in t for stop in BODY_STOP_TOKENS):
            break
    return None

# -------------------------------------------
# API principal
# -------------------------------------------
def insert_logo(docx_bytes: bytes,
                image_bytes: bytes,
                *,
                max_width_cm: float = 10.0,
                prefer_header: bool = True,
                center: bool = True,
                pad_cm: float = 0.0) -> bytes:
    """
    Substitui o placeholder 'Logo da empresa' pela imagem (inline) na PRIMEIRA PÁGINA.

    Parâmetros:
      - docx_bytes: bytes do documento Word base.
      - image_bytes: bytes da imagem (PNG/JPG etc).
      - max_width_cm: largura máxima do logotipo (escala proporcional). Default: 10 cm.
      - prefer_header: se True, tenta substituir no cabeçalho da 1ª seção primeiro.
      - center: se True, centraliza o parágrafo do logo.
      - pad_cm: reserva adicional subtraída da largura útil da página (se quiser margem extra).

    Retorna:
      - bytes do DOCX com o logo inserido; se algo falhar, devolve o original.
    """
    try:
        base = bytes(docx_bytes) if not isinstance(docx_bytes, (bytes, bytearray)) else docx_bytes
        doc = Document(BytesIO(base))

        # 1) Determina alvo: header > body (1ª página)
        target_p: Optional[Paragraph] = None
        if prefer_header:
            target_p = _find_placeholder_in_pars(_all_header_paragraphs(doc))
        if target_p is None:
            target_p = _find_placeholder_in_body_first_page(doc)
        if target_p is None:
            # nada a fazer; devolve original
            return base

        # 2) Limpa o parágrafo e insere a imagem
        if center:
            try:
                target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        _clear_runs(target_p)
        # largura final: min(max_width_cm, largura útil)
        width_cm = min(max_width_cm, _available_width_cm(doc, pad_cm=pad_cm))
        run = target_p.add_run()
        run.add_picture(BytesIO(image_bytes), width=Cm(width_cm))

        # força apenas cor preta dos runs subsequentes (se houver)
        _force_following_runs_black(target_p, from_index=0)

        # 3) Salva
        out = BytesIO()
        doc.save(out); out.seek(0)
        final = out.read()
        if not isinstance(final, (bytes, bytearray)) or len(final) == 0:
            return base
        return final
    except Exception:
        # falha segura → devolve original
        try:
            return bytes(docx_bytes)
        except Exception:
            return b""