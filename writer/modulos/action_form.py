# modulos/action_form.py
from __future__ import annotations
from io import BytesIO
from typing import Dict, Optional, List, Tuple
import re, unicodedata
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

# ============================================================
# Normalização / util
# ============================================================
def _norm_ascii_lower(s: str) -> str:
    # NBSP (Word) -> espaço normal
    s = (s or "").replace("\u00A0", " ")
    return "".join(
        c for c in unicodedata.normalize("NFKD", s)
        if not unicodedata.combining(c)
    ).lower()

def _para_text(p: Paragraph) -> str:
    return p.text or ""

def _is_effectively_empty(p: Paragraph) -> bool:
    return (p.text or "").strip() == ""

def _split_to_template_paragraphs(value: str) -> List[str]:
    """
    Divide o texto em parágrafos lógicos por dupla quebra (\n\n).
    Dentro de cada parágrafo, converte quebras simples em espaço e comprime espaços.
    NÃO cria <w:br/>; cada item é mapeado a um parágrafo placeholder existente.
    """
    if not value:
        return []
    v = value.replace("\r", "\n").replace("\t", " ").replace("\u00A0", " ")
    blocks = re.split(r"\n\s*\n", v)
    out: List[str] = []
    for blk in blocks:
        txt = re.sub(r"\s*\n\s*", " ", blk.strip())
        txt = re.sub(r"[ ]{2,}", " ", txt).strip()
        if txt:
            out.append(txt)
    return out

# ============================================================
# Coleta de parágrafos (corpo + tabelas)
# ============================================================
def _all_paragraphs(doc: Document) -> List[Paragraph]:
    pars = list(doc.paragraphs)  # corpo
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                pars.extend(cell.paragraphs)
    return pars

# ============================================================
# Janela da seção
# ============================================================
START_TOKENS = ("programa e atividade de pd&i", "programa/atividade de pd&i")
END_TOKENS = (
    "dispêndios","dispendios","dispêndios do programa","dispendios do programa",
    "incentivos fiscais do programa","observações gerais","observacoes gerais"
)

def _find_section_span(pars: List[Paragraph]) -> Tuple[int, int]:
    texts = [_norm_ascii_lower(p.text) for p in pars]
    i_start = next((i for i, t in enumerate(texts) if any(tok in t for tok in START_TOKENS)), -1)
    if i_start == -1:
        return -1, -1
    i_end = next((j for j in range(i_start + 1, len(texts)) if any(tk in texts[j] for tk in END_TOKENS)), len(texts))
    return i_start, i_end

# ============================================================
# Labels (regex) e mapeamento
# ============================================================
LAB_DESTAQUE = re.compile(r"\bdestaque\b.*\belemento\b.*\btecnolog|inovador", re.IGNORECASE)
LAB_BARREIRA = re.compile(r"\b(barreira|desafio)\b.*\btecnolog", re.IGNORECASE)
LAB_METODO   = re.compile(r"\bmetodolog(?:ia)?(?:\s*/\s*m[eé]todos)?", re.IGNORECASE)

LAB_PBPADE   = re.compile(r"\bpb,\s*pa\s*e\s*de\b", re.IGNORECASE)
LAB_NATUREZA = re.compile(r"\bnatureza\b.*\bprojeto.*formpd\b", re.IGNORECASE)
LAB_ATIV_CONT= re.compile(r"\batividade\b.*\bcont[ií]nua", re.IGNORECASE)

LAB_DATA_INI = re.compile(r"\bdata\b.*\bin[ií]cio\b.*\batividade", re.IGNORECASE)
LAB_PREV_TERM= re.compile(r"\bprevis[aã]o\b.*\bt[ée]rmino", re.IGNORECASE)

# Identificação (inline)
LAB_NOME_PROJETO   = re.compile(r"(?:\b\d+\s*[\.\)]?\s*)?\bnome\s+do\s+projeto\b", re.IGNORECASE)
LAB_NUMERO_PROJETO = re.compile(r"\bn[úu]mero\s+do\s+projeto\b", re.IGNORECASE)

# Nome da atividade de PD&I (inline) — tolerante a PD&I / PD I / P-D-I / P D & I
LAB_NOME_ATIV_PDI  = re.compile(r"\bnome\s+da\s+atividade\s+de\s+(?:pd\W*&?\W*i|pdi)\b", re.IGNORECASE)

FIELD_MAP = [
    # === Textuais: valor na linha seguinte (next_single) ===
    ("destaque_elemento_novo",       LAB_DESTAQUE, "next_single"),
    ("barreira_desafio_tecnologico", LAB_BARREIRA, "next_single"),
    ("metodologia_metodos",          LAB_METODO,   "next_single"),

    # Inline (como já funcionava)
    ("pb_pa_ou_de",                  LAB_PBPADE,   "inline"),
    ("natureza",                     LAB_NATUREZA, "inline"),

    # 🔁 TROCA SOLICITADA: 'atividade_continua' agora é next_single
    ("atividade_continua",           LAB_ATIV_CONT,"next_single"),

    # Identificação (inline)
    ("nome_projeto",                 LAB_NOME_PROJETO,   "inline"),
    ("numero_projeto",               LAB_NUMERO_PROJETO, "inline"),

    # Inline do campo problemático (com escritor dedicado)
    ("nome_atividade_pdi",           LAB_NOME_ATIV_PDI,  "inline"),

    # Linha seguinte única (curtos)
    ("data_inicio_atividade",        LAB_DATA_INI, "next_single"),
    ("previsao_termino",             LAB_PREV_TERM,"next_single"),
]

STOP_LABELS_FOR_BLOCK = [
    LAB_DESTAQUE, LAB_BARREIRA, LAB_METODO,
    LAB_PBPADE, LAB_NATUREZA, LAB_ATIV_CONT, LAB_DATA_INI, LAB_PREV_TERM,
    LAB_NOME_PROJETO, LAB_NUMERO_PROJETO, LAB_NOME_ATIV_PDI
]

# ============================================================
# Auxiliares de escrita / bloco
# ============================================================
def _index_of_next_label(pars: List[Paragraph], start_idx: int, end_idx: int) -> int:
    for i in range(start_idx, end_idx):
        if any(pat.search(_norm_ascii_lower(_para_text(pars[i]))) for pat in STOP_LABELS_FOR_BLOCK):
            return i
    return end_idx

def _collect_placeholder_block(pars: List[Paragraph], label_idx: int, end_idx: int) -> List[int]:
    start = label_idx + 1
    if start >= end_idx or start >= len(pars):
        return []
    stop = _index_of_next_label(pars, start, end_idx)
    return list(range(start, max(start, stop)))

def _overwrite_paragraph_text_keep_runs(p: Paragraph, new_text: str):
    if not p.runs:
        p.add_run("")  # ancora
    for r in p.runs:
        r.text = ""
    p.runs[0].text = new_text or ""

def _delete_paragraph(p: Paragraph):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

def _collapse_trailing_empty_in_block(pars: List[Paragraph], block_idxs: List[int], keep: int = 1) -> bool:
    if not block_idxs:
        return False
    trailing = 0
    for idx in reversed(block_idxs):
        if _is_effectively_empty(pars[idx]):
            trailing += 1
        else:
            break
    removed = False
    if trailing > keep:
        to_remove = trailing - keep
        i = len(block_idxs) - 1
        while to_remove > 0 and i >= 0:
            idx = block_idxs[i]
            if _is_effectively_empty(pars[idx]):
                _delete_paragraph(pars[idx])
                removed = True
                to_remove -= 1
            i -= 1
    return removed

def _overwrite_block_with_value(doc: Document, pars: List[Paragraph], block_idxs: List[int], value: str) -> bool:
    segs = _split_to_template_paragraphs(value)
    if not block_idxs:
        return False
    if not segs:
        for idx in block_idxs:
            _overwrite_paragraph_text_keep_runs(pars[idx], "")
        return _collapse_trailing_empty_in_block(pars, block_idxs, keep=1)

    n, m = len(block_idxs), len(segs)
    if m <= n:
        for i in range(m):
            _overwrite_paragraph_text_keep_runs(pars[block_idxs[i]], segs[i])
        for j in range(m, n):
            _overwrite_paragraph_text_keep_runs(pars[block_idxs[j]], "")
        return _collapse_trailing_empty_in_block(pars, block_idxs, keep=1)
    else:
        for i in range(n - 1):
            _overwrite_paragraph_text_keep_runs(pars[block_idxs[i]], segs[i])
        remainder = " ".join(segs[n - 1:])
        _overwrite_paragraph_text_keep_runs(pars[block_idxs[-1]], remainder)
        return False

# --------- ajustes de cor ---------
def _normalize_run_to_black(run):
    """
    Força PRETO retirando ênfases (uso: 'natureza', como no fluxo original).
    """
    f = run.font
    f.bold = False; f.italic = False; f.underline = None
    try: f.highlight_color = None
    except Exception: pass
    try: f.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception: pass

def _force_run_color_black(run):
    """Força apenas a cor preta do valor (mantém negrito/itálico do placeholder)."""
    try:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception:
        pass

# --------- inline padrão ---------
def _fill_inline_after_colon_simple(
    p: Paragraph, value: str, *, force_black: bool = False, force_color_black_only: bool = False
):
    """
    Escreve o valor APÓS o primeiro ":" no mesmo parágrafo (inline),
    reaproveitando o run imediatamente após ":" e limpando os demais (substituição real).
    """
    runs = p.runs
    if not runs:
        r = p.add_run(": " + (value or ""))
        if force_black: _normalize_run_to_black(r)
        elif force_color_black_only: _force_run_color_black(r)
        return

    # localizar ':'
    col_i, col_pos = -1, -1
    for i, r in enumerate(runs):
        t = r.text or ""
        pos = t.find(":")
        if pos >= 0:
            col_i, col_pos = i, pos
            break

    if col_i == -1:
        runs[-1].text = (runs[-1].text or "").rstrip() + ":"
        r = p.add_run(" " + (value or ""))
        if force_black: _normalize_run_to_black(r)
        elif force_color_black_only: _force_run_color_black(r)
        return

    # garante "Rótulo:"
    runs[col_i].text = (runs[col_i].text or "")[:col_pos+1]

    # valor no run imediatamente após ':', limpando os demais
    if col_i + 1 < len(runs):
        runs[col_i + 1].text = " " + (value or "")
        for k in range(col_i + 2, len(runs)):
            runs[k].text = ""
        if force_black:
            _normalize_run_to_black(runs[col_i + 1])
        elif force_color_black_only:
            _force_run_color_black(runs[col_i + 1])
    else:
        r = p.add_run(" " + (value or ""))
        if force_black:
            _normalize_run_to_black(r)
        elif force_color_black_only:
            _force_run_color_black(r)

# --------- inline "smart" EXCLUSIVO para 'Nome da atividade de PD&I' ---------
def _fill_inline_after_colon_smart_nome_atividade(p: Paragraph, value: str):
    """
    Para o rótulo 'Nome da atividade de PD&I':
    - acha o ':' em qualquer run,
    - trunca o run no ':',
    - apaga TODO conteúdo à direita (inclui textos-modelo como '... / Sistema de contestação'),
    - escreve o valor no 1º run após o ':', forçando só a cor preta do valor (mantém negrito).
    """
    runs = p.runs
    if not runs:
        r = p.add_run(": " + (value or ""))
        _force_run_color_black(r)
        return

    col_i, col_pos = -1, -1
    for i, r in enumerate(runs):
        t = r.text or ""
        pos = t.find(":")
        if pos >= 0:
            col_i, col_pos = i, pos
            break

    if col_i == -1:
        runs[-1].text = (runs[-1].text or "").rstrip() + ":"
        r = p.add_run(" " + (value or ""))
        _force_run_color_black(r)
        return

    # corta no ':' e apaga à direita
    runs[col_i].text = (runs[col_i].text or "")[:col_pos + 1]
    if col_i + 1 < len(runs):
        runs[col_i + 1].text = " " + (value or "")
        for k in range(col_i + 2, len(runs)):
            runs[k].text = ""
        _force_run_color_black(runs[col_i + 1])
    else:
        r = p.add_run(" " + (value or ""))
        _force_run_color_black(r)

# ============================================================
# Preenchimento (com heurística de rótulo curto para textuais)
# ============================================================
def _fill_fields_in_section(doc: Document, fields: Dict[str, Optional[str]]):
    pars = _all_paragraphs(doc)
    i_start, i_end = _find_section_span(pars)
    if i_start == -1:
        return

    # Rótulos textuais sensíveis → precisam ter ':' E ser "curtos"
    SENSITIVE_KEYS = {
        "destaque_elemento_novo",
        "barreira_desafio_tecnologico",
        "metodologia_metodos",
    }
    MAX_LABEL_LEN = 160  # rótulo curto; conteúdo tende a ser bem maior

    def _is_label_paragraph(key: str, raw: str, txt_norm: str) -> bool:
        if key in SENSITIVE_KEYS:
            if ":" not in raw:
                return False
            if len(txt_norm) > MAX_LABEL_LEN:
                return False
            if key == "destaque_elemento_novo":
                return ("destaque" in txt_norm and "elemento" in txt_norm and ("tecnolog" in txt_norm or "inovador" in txt_norm))
            if key == "barreira_desafio_tecnologico":
                return (("barreira" in txt_norm or "desafio" in txt_norm) and "tecnolog" in txt_norm)
            if key == "metodologia_metodos":
                return ("metodolog" in txt_norm and ("utilizad" in txt_norm or "metodo" in txt_norm))
        return True

    for key, pat, mode in FIELD_MAP:
        val = (fields.get(key) or "").strip()

        # Fallbacks pontuais
        if key == "numero_projeto" and not val:
            val = "1"  # padrão
        if key == "nome_atividade_pdi" and not val:
            # se não vier, usa o mesmo nome do projeto
            val = (fields.get("nome_projeto") or "").strip()

        if not val:
            continue

        # -------- localizar o rótulo (regex + guarda + heurística) --------
        found = -1
        for i in range(i_start, min(i_end, len(pars))):
            raw = _para_text(pars[i])
            txt = _norm_ascii_lower(raw)
            if pat.search(txt) and _is_label_paragraph(key, raw, txt):
                found = i
                break

        if found == -1:
            continue

        try:
            if key == "nome_atividade_pdi":
                _fill_inline_after_colon_smart_nome_atividade(pars[found], val)
            elif mode == "inline":
                _fill_inline_after_colon_simple(
                    pars[found], val,
                    force_black=(key == "natureza"),
                    force_color_black_only=(key in ("nome_projeto", "numero_projeto"))
                )
            elif mode == "next_single":
                j = found + 1
                if j < len(pars):
                    segs = _split_to_template_paragraphs(val)
                    text_single = " ".join(segs) if segs else ""
                    _overwrite_paragraph_text_keep_runs(pars[j], text_single)
                else:
                    _fill_inline_after_colon_simple(pars[found], val)
            elif mode == "block_overwrite":
                block_idxs = _collect_placeholder_block(pars, found, i_end)
                removed = _overwrite_block_with_value(doc, pars, block_idxs, val)
                if removed:
                    pars = _all_paragraphs(doc)
                    i_start, i_end = _find_section_span(pars)

            if key == "previsao_termino":
                break

        except Exception:
            # resiliência: um erro não trava os demais
            continue

# ============================================================
# API pública
# ============================================================
def fill_formpd_section(template_bytes: bytes, form_fields: Dict[str, Optional[str]]) -> bytes:
    """
    Abre o .docx (bytes), preenche a seção 'Programa e Atividade de PD&I'
    e retorna o .docx atualizado (bytes). Em erro, devolve o original.
    """
    try:
        base = bytes(template_bytes) if not isinstance(template_bytes, (bytes, bytearray)) else template_bytes
        doc = Document(BytesIO(base))
        _fill_fields_in_section(doc, form_fields or {})
        out = BytesIO(); doc.save(out); out.seek(0)
        final = out.read()
        if not isinstance(final, (bytes, bytearray)) or len(final) == 0:
            return bytes(base)
        return bytes(final)
    except Exception:
        try:
            return bytes(template_bytes)
        except Exception:
            return b""
