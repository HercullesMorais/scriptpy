# modulos/action_rh.py
from __future__ import annotations
import re, unicodedata
from io import BytesIO
from typing import Any, Dict, List, Tuple, Optional
from copy import deepcopy

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

__all__ = [
    "locate_relacao_rh",
    "fill_exclusives_from_sheet",
    "fill_exclusive_from_sheet",  # alias
    "parse_rh_sheet",             # novo
    "fill_exclusives_from_items", # novo (rápido)
]

# ---------------------------- regex pré-compiladas ----------------------------
PAT_NUM_ITEM = re.compile(r"\bnumero do item\b|\bnúmero do item\b", re.IGNORECASE)
PAT_CPF      = re.compile(r"\bcpf\b", re.IGNORECASE)
PAT_NOME     = re.compile(r"\bnome\b", re.IGNORECASE)
PAT_TIT      = re.compile(r"titulacao|titulação", re.IGNORECASE)
PAT_HORAS    = re.compile(r"total de horas \(anual\)|total horas \(anual\)", re.IGNORECASE)
PAT_DED      = re.compile(r"dedicacao|dedicação", re.IGNORECASE)
PAT_DECL     = re.compile(r"valor total declarado \(r\$\)", re.IGNORECASE)
PAT_APR_P    = re.compile(r"valor total aprovado parecer \(r\$\)", re.IGNORECASE)
PAT_APR_C    = re.compile(r"valor total aprovado na contestacao \(r\$\)|valor total aprovado na contestação \(r\$\)", re.IGNORECASE)
PAT_SOLIC    = re.compile(r"total solicitado no recurso administrativo \(r\$\)", re.IGNORECASE)
PAT_LABEL_J  = re.compile(r"solicitacao recurso administrativo|solicitação recurso administrativo", re.IGNORECASE)

TITLE_RX     = re.compile(r"rela[cç][aã]o de recursos humanos", re.IGNORECASE)
EXCL_RX      = re.compile(r"recursos humanos exclusivos", re.IGNORECASE)
PARC_RX      = re.compile(r"recursos humanos parciais", re.IGNORECASE)

# -------------------------- utils de texto/docx ---------------------------

def _norm(s: str) -> str:
    s = s or ""
    s = "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s


def _all_paragraphs(doc: Document) -> List[Paragraph]:
    pars: List[Paragraph] = []
    pars.extend(doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                pars.extend(cell.paragraphs)
    for section in doc.sections:
        if section.header:
            pars.extend(section.header.paragraphs)
            for t in section.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        pars.extend(c.paragraphs)
        if section.footer:
            pars.extend(section.footer.paragraphs)
            for t in section.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        pars.extend(c.paragraphs)
    return pars


def _is_nan(v: Any) -> bool:
    try:
        if pd is not None:
            return bool(pd.isna(v))
    except Exception:
        pass
    return False


def _clean_str(v: Any) -> str:
    if v is None or _is_nan(v):
        return ""
    return str(v)


def _replace_after_colon(p: Paragraph, value: str) -> bool:
    runs = p.runs
    if not runs:
        p.add_run(": " + value)
        return True
    full = "".join(r.text or "" for r in runs)
    pos = full.find(":")
    if pos < 0:
        if runs:
            runs[-1].text = (runs[-1].text or "").rstrip() + ": "
        else:
            p.add_run(": ")
        p.add_run(value)
        return True
    count = 0
    for i, r in enumerate(runs):
        t = r.text or ""
        if count + len(t) <= pos:
            count += len(t)
            continue
        offset = pos - count
        if offset >= 0:
            r.text = t[: offset + 1]
            if i + 1 < len(runs):
                runs[i + 1].text = " " + value
                for k in range(i + 2, len(runs)):
                    runs[k].text = ""
            else:
                p.add_run(" " + value)
            return True
        else:
            r.text = ""
    p.add_run(" " + value)
    return True


def _replace_next_paragraph_contentonly(p_label: Paragraph, following: List[Paragraph], value: str) -> bool:
    for p_next in following:
        txt = (p_next.text or "").strip()
        if not txt or (":" not in txt):
            if not p_next.runs:
                p_next.add_run("")
            for i in range(1, len(p_next.runs)):
                p_next.runs[i].text = ""
            p_next.runs[0].text = value
            return True
    return _replace_after_colon(p_label, value)


def _set_paragraph_text(p: Paragraph, text: str) -> None:
    if not p.runs:
        p.add_run("")
    for i in range(1, len(p.runs)):
        p.runs[i].text = ""
    p.runs[0].text = text

# -------------------------- localização de seções -------------------------

def _find_header_index(texts_norm: List[str], pat: re.Pattern, start: int = 0, end: Optional[int] = None) -> int:
    end = len(texts_norm) if end is None else end
    for i in range(start, end):
        if pat.search(texts_norm[i]):
            return i
    return -1


def _locate_on_doc(doc: Document, window_fwd_scan: int = 2000) -> Tuple[bool, Dict[str, Any]]:
    pars = _all_paragraphs(doc)
    texts_norm = [_norm(p.text) for p in pars]
    title_idx = _find_header_index(texts_norm, TITLE_RX, 0, len(texts_norm))
    if title_idx == -1:
        return False, {"reason": "target_not_found"}

    excl_idx = _find_header_index(texts_norm, EXCL_RX, title_idx, min(len(texts_norm), title_idx + window_fwd_scan))
    parc_idx = _find_header_index(texts_norm, PARC_RX, title_idx, min(len(texts_norm), title_idx + window_fwd_scan))

    excl_j0 = excl_idx + 1 if excl_idx != -1 else -1
    excl_j1 = parc_idx if (parc_idx != -1 and excl_idx != -1) else min(len(texts_norm), (excl_j0 if excl_j0!=-1 else title_idx) + 800)

    parc_j0 = parc_idx + 1 if parc_idx != -1 else -1
    parc_j1 = min(len(texts_norm), parc_j0 + 800) if parc_j0 != -1 else -1

    return True, {
        "exclusivos_found": excl_idx != -1,
        "exclusivos_index": excl_idx,
        "parciais_found": parc_idx != -1,
        "parciais_index": parc_idx,
        "search_window_exclusivos": (excl_j0, excl_j1) if excl_j0 != -1 else None,
        "search_window_parciais": (parc_j0, parc_j1) if parc_j0 != -1 else None,
    }


def locate_relacao_rh(docx_bytes: bytes,
                      window_back_disp: int = 200,
                      window_fwd_scan: int = 2000) -> Tuple[bool, Dict[str, Any]]:
    doc = Document(BytesIO(docx_bytes))
    return _locate_on_doc(doc, window_fwd_scan)

# ------------------------- leitura/formatos planilha ------------------

def _fmt_moeda_br(v: Any) -> str:
    if isinstance(v, str):
        s = v.strip()
        if s and not _is_nan(s):
            return s
    try:
        x = float(v)
    except Exception:
        return ""
    s = f"{x:,.2f}".replace(",","X").replace(".",",").replace("X",".")
    return s


def _fmt_horas(v: Any) -> str:
    try:
        x = float(v)
        if abs(x - int(x)) < 1e-9:
            return str(int(x))
        return f"{x:.1f}".replace(".", ",")
    except Exception:
        return _clean_str(v)


def _pick_value(row: Dict[str, Any], *cands: Optional[str]) -> Optional[Any]:
    for c in cands:
        if not c: continue
        if c in row and row[c] not in (None, "") and not _is_nan(row[c]):
            return row[c]
    return None

# --- NEW: parse sheet (cache no app) ---

def parse_rh_sheet(xlsx_bytes: bytes, sheet_name: str | None = None) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    if pd is None:
        raise RuntimeError("pandas não disponível para leitura do XLSX")
    with BytesIO(xlsx_bytes) as bio:
        df = pd.read_excel(bio, engine="openpyxl", sheet_name=sheet_name or 0)
    lower = {str(c).strip().lower(): c for c in df.columns}
    ded_col = lower.get("dedicação") or lower.get("dedicacao")
    if not ded_col:
        return [], []

    cand_val_decl = (
        lower.get("total declarado (r$)"),
        lower.get("valor total declarado (r$)"),
        lower.get("valor (r$)"),
        lower.get("total solicitado no recurso administrativo (r$)"),
    )

    def _row_to_item(row) -> Dict[str, Any]:
        return {
            "Item": _pick_value(row, lower.get("item"), "Item"),
            "CPF": _pick_value(row, lower.get("cpf"), "CPF"),
            "Nome": _pick_value(row, lower.get("nome"), "Nome"),
            "Titulação": _pick_value(row, lower.get("titulação"), lower.get("titulacao"), "Titulação", "Titulacao"),
            "Total horas (Anual)": _pick_value(row, lower.get("total horas (anual)"), lower.get("total de horas"), "Total horas (Anual)", "Total de horas"),
            "Dedicação": _pick_value(row, ded_col, "Dedicação", "Dedicacao"),
            "Valor Declarado": _pick_value(row, *cand_val_decl),
            "Justificativa": _pick_value(row, lower.get("justificativa"), "Justificativa"),
        }

    excl_items, parc_items = [], []
    for _, s in df.iterrows():
        row = s.to_dict()
        dedic = _norm(str(row.get(ded_col) or ""))
        item = _row_to_item(row)
        if "exclusiv" in dedic:
            excl_items.append(item)
        elif "parcia" in dedic:
            parc_items.append(item)
    return excl_items, parc_items

# --------------------- indexação do bloco (uma vez) ----------------------

def _index_labels(block_texts: List[str]) -> Dict[str, int]:
    idx = {}
    for i, t in enumerate(block_texts):
        if 'num_item' not in idx and PAT_NUM_ITEM.search(t): idx['num_item'] = i
        elif 'cpf' not in idx and PAT_CPF.search(t): idx['cpf'] = i
        elif 'nome' not in idx and PAT_NOME.search(t): idx['nome'] = i
        elif 'tit' not in idx and PAT_TIT.search(t): idx['tit'] = i
        elif 'horas' not in idx and PAT_HORAS.search(t): idx['horas'] = i
        elif 'ded' not in idx and PAT_DED.search(t): idx['ded'] = i
        elif 'decl' not in idx and PAT_DECL.search(t): idx['decl'] = i
        elif 'aprp' not in idx and PAT_APR_P.search(t): idx['aprp'] = i
        elif 'aprc' not in idx and PAT_APR_C.search(t): idx['aprc'] = i
        elif 'sol' not in idx and PAT_SOLIC.search(t): idx['sol'] = i
        elif 'labelj' not in idx and PAT_LABEL_J.search(t): idx['labelj'] = i
    if 'labelj' in idx:
        idx['just'] = idx['labelj'] + 1
    return idx

# ------------------------- localizar bloco modelo ------------------

def _find_model_block(pars: List[Paragraph], texts_norm: List[str], j0: int, j1: int) -> Tuple[int, int, Dict[str,int]]:
    start = -1
    for i in range(j0, j1):
        if PAT_NUM_ITEM.search(texts_norm[i]):
            start = i; break
    if start == -1:
        return -1, -1, {}
    label = -1
    for i in range(start, j1):
        if PAT_LABEL_J.search(texts_norm[i]):
            label = i; break
    end = (label + 1) if (label != -1 and (label + 1) < j1) else min(start + 12, j1-1)
    block_texts = [texts_norm[i] for i in range(start, end+1)]
    label_idx = _index_labels(block_texts)
    return start, end, label_idx


def _insert_blank_after(par: Paragraph) -> Paragraph:
    new_p = OxmlElement('w:p')
    par._p.addnext(new_p)
    return Paragraph(new_p, par._parent)


def _duplicate_block_after(doc: Document, block_pars: List[Paragraph], copies: int) -> List[List[Paragraph]]:
    if copies <= 0:
        return []
    result: List[List[Paragraph]] = []
    anchor = block_pars[-1]
    # quebra de linha após bloco original para separar do próximo
    anchor = _insert_blank_after(anchor)
    for i in range(copies):
        new_block: List[Paragraph] = []
        for p in block_pars:
            new_p = deepcopy(p._p)
            anchor._p.addnext(new_p)
            anchor = Paragraph(new_p, anchor._parent)
            new_block.append(anchor)
        result.append(new_block)
        # quebra de linha entre blocos clonados
        if i < copies - 1:
            anchor = _insert_blank_after(anchor)
    return result

# ------------------------- preenchimento rápido ------------------

def _fill_block_by_index(pars: List[Paragraph], base: int, label_idx: Dict[str,int], item: Dict[str, Any]):
    def _put(key: str, value: str):
        if key in label_idx:
            _replace_after_colon(pars[base + label_idx[key]], value)
    _put('num_item', _clean_str(item.get('Item')).strip())
    _put('cpf', _clean_str(item.get('CPF')).strip())
    _put('nome', _clean_str(item.get('Nome')).strip())
    _put('tit', _clean_str(item.get('Titulação')).strip())
    _put('horas', _fmt_horas(item.get('Total horas (Anual)')))
    _put('ded', _clean_str(item.get('Dedicação') or ''))
    _put('decl', _fmt_moeda_br(item.get('Valor Declarado')))
    _put('aprp', '0')
    _put('aprc', '0')
    _put('sol',  _fmt_moeda_br(item.get('Valor Declarado')))
    just_text = _clean_str(item.get('Justificativa')).strip()
    if just_text:
        if 'just' in label_idx:
            _set_paragraph_text(pars[base + label_idx['just']], just_text)
        elif 'labelj' in label_idx:
            jpos = base + label_idx['labelj'] + 1
            _replace_next_paragraph_contentonly(pars[base + label_idx['labelj']], pars[jpos:jpos+4], just_text)

# ------------------------- pipelines ------------------

def _process_category_fast(doc: Document,
                           window: Tuple[int,int],
                           items: List[Dict[str, Any]],
                           dedic_default: str) -> Tuple[int, List[Tuple[int,int]]]:
    if not items or not window:
        return 0, []
    pars = _all_paragraphs(doc)
    texts_norm = [_norm(p.text) for p in pars]
    j0, j1 = window
    if j0 < 0 or j1 <= j0:
        return 0, []

    b0, b1, label_idx = _find_model_block(pars, texts_norm, j0, j1)
    if b0 == -1 or not label_idx:
        return 0, []

    model_block = [pars[i] for i in range(b0, b1+1)]
    copies = _duplicate_block_after(doc, model_block, max(0, len(items)-1))

    # recompute paragraph refs post clone
    pars = _all_paragraphs(doc)

    ranges: List[Tuple[int,int]] = [(b0, b1)]
    for blk in copies:
        start_idx = None; end_idx = None
        for k, p in enumerate(pars):
            if p._p is blk[0]._p: start_idx = k
            if p._p is blk[-1]._p: end_idx = k
            if start_idx is not None and end_idx is not None: break
        if start_idx is not None and end_idx is not None:
            ranges.append((start_idx, end_idx))

    processed = 0
    for i, (s,e) in enumerate(ranges):
        if i >= len(items): break
        if not items[i].get('Dedicação'): items[i]['Dedicação'] = dedic_default
        _fill_block_by_index(pars, s, label_idx, items[i])
        processed += 1
    return processed, ranges

# Public: rápido sem pandas na geração

def fill_exclusives_from_items(docx_bytes: bytes,
                               exclusives: List[Dict[str, Any]] | None,
                               parciais: List[Dict[str, Any]] | None,
                               window_fwd_scan: int = 2000) -> Tuple[bytes, Dict[str, Any]]:
    doc = Document(BytesIO(docx_bytes))

    ok0, meta0 = _locate_on_doc(doc, window_fwd_scan)
    if not ok0:
        out = BytesIO(); doc.save(out); out.seek(0)
        return out.read(), {"reason": "locate_fail"}

    processed_excl, ranges_excl = 0, []
    if exclusives and meta0.get('search_window_exclusivos'):
        processed_excl, ranges_excl = _process_category_fast(doc, meta0['search_window_exclusivos'], exclusives, dedic_default='Exclusivo')

    # break -> relocaliza parciais no doc atualizado (sem IO)
    ok1, meta1 = _locate_on_doc(doc, window_fwd_scan)

    processed_parc, ranges_parc = 0, []
    if parciais and meta1.get('search_window_parciais'):
        processed_parc, ranges_parc = _process_category_fast(doc, meta1['search_window_parciais'], parciais, dedic_default='Parcial')

    out = BytesIO(); doc.save(out); out.seek(0)
    return out.read(), {
        "processed_exclusivos": processed_excl,
        "processed_parciais": processed_parc,
        "ranges_exclusivos": ranges_excl,
        "ranges_parciais": ranges_parc,
    }

# Mantém compat de bytes (lento, pois lê XLSX):

def fill_exclusives_from_sheet(docx_bytes: bytes,
                               xlsx_bytes: bytes,
                               sheet_name: str | None = None,
                               search_window_ahead: int = 2000) -> Tuple[bytes, Dict[str, Any]]:
    excl, parc = parse_rh_sheet(xlsx_bytes, sheet_name)
    return fill_exclusives_from_items(docx_bytes, excl, parc, search_window_ahead)

# alias compat

def fill_exclusive_from_sheet(docx_bytes: bytes,
                              xlsx_bytes: bytes,
                              sheet_name: str | None = None,
                              search_window_ahead: int = 2000) -> Tuple[bytes, Dict[str, Any]]:
    return fill_exclusives_from_sheet(docx_bytes, xlsx_bytes, sheet_name, search_window_ahead)
