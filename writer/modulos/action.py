# modulos/action.py
from typing import Dict, Optional, List, Tuple
from io import BytesIO
from datetime import date
import re, unicodedata

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P

# --- UF -> Estado (compacto) ---
_pairs = "AC:Acre;AL:Alagoas;AP:Amapá;AM:Amazonas;BA:Bahia;CE:Ceará;DF:Distrito Federal;ES:Espírito Santo;GO:Goiás;MA:Maranhão;MT:Mato Grosso;MS:Mato Grosso do Sul;MG:Minas Gerais;PA:Pará;PB:Paraíba;PR:Paraná;PE:Pernambuco;PI:Piauí;RJ:Rio de Janeiro;RN:Rio Grande do Norte;RS:Rio Grande do Sul;RO:Rondônia;RR:Roraima;SC:Santa Catarina;SP:São Paulo;SE:Sergipe;TO:Tocantins"
UF_TO_ESTADO = {k: v for k, v in (p.split(":") for p in _pairs.split(";"))}

# --- helpers de formatação/comparação (enxutos) ---
def _fmt_cnpj(cnpj: Optional[str]) -> str:
    if not cnpj: return ""
    d = re.sub(r"\D", "", cnpj)
    return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}" if len(d)==14 else cnpj.strip()

def _fmt_cep(cep: Optional[str]) -> Tuple[str, str]:
    if not cep: return "", ""
    d = re.sub(r"\D", "", cep)
    return (d[:5], d[5:]) if len(d)==8 else (cep.replace("-", "").strip(), "")

def _mes_nome_pt(m: int) -> str:
    nomes = ["janeiro","fevereiro","março","abril","maio","junho","julho","agosto","setembro","outubro","novembro","dezembro"]
    return nomes[m-1] if 1<=m<=12 else "mês"

def _norm_ascii_lower(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)).lower()

# --- contexto (mesmos campos) ---
def build_capa_context(
    *, empresa: Optional[str], cnpj: Optional[str], parecer_contestacao: Optional[str],
    ano_base: Optional[int], logradouro: Optional[str], numero: Optional[str],
    cep: Optional[str], municipio: Optional[str], uf: Optional[str],
    lote: Optional[str]=None, estado_nome: Optional[str]=None,
    dia: Optional[int]=None, mes: Optional[int]=None, ano_data: Optional[int]=None,
    responsavel: Optional[str]=None, parecer_tecnico: Optional[str]=None
) -> Dict[str,str]:
    today = date.today()
    dia = int(dia or today.day); mes = int(mes or today.month); ano_data = int(ano_data or today.year)
    uf_norm = (uf or "").strip().upper() or None
    estado_ok = (estado_nome or UF_TO_ESTADO.get(uf_norm, "")).strip()
    cnpj_ok = _fmt_cnpj(cnpj); cep5, cep3 = _fmt_cep(cep)
    ab_full = str(ano_base or ano_data)
    return {
        "empresa": (empresa or "").strip(),
        "cnpj": cnpj_ok,
        "parecer": (parecer_contestacao or "").strip(),
        "parecer_tecnico": (parecer_tecnico or "").strip(),
        "ano_base_full": ab_full,
        "ano_base_2": ab_full[-2:],
        "logradouro": (logradouro or "").strip(),
        "numero": (numero or "").strip(),
        "cep_5": cep5, "cep_3": cep3,
        "municipio": (municipio or "").strip(),
        "uf": (uf or "").strip().upper(),
        "lote": (lote or "").strip(),
        "estado_nome": estado_ok,
        "dia2": f"{dia:02d}",
        "mes_ext": _mes_nome_pt(mes),
        "ano_data_2": str(ano_data)[-2:],
        "responsavel": (responsavel or "").strip(),
    }

# --- parágrafos até "Sumário" ---
def _iter_paragraphs_before_sumario(doc: Document) -> List[Paragraph]:
    pars: List[Paragraph] = []
    for el in doc.element.body.iter():
        if isinstance(el, CT_P) or str(getattr(el,"tag","")).endswith("}p"):
            p = Paragraph(el, doc)
            if p.text.strip().startswith("Sumário"): break
            pars.append(p)
    return pars

# --- busca em runs (curtas) ---
def _first_run_eq(runs, t: str, start=0) -> int:
    for i in range(start, len(runs)):
        if runs[i].text == t: return i
    return -1

def _find_pair(runs, a: str, b: str, start=0) -> int:
    for i in range(start, len(runs)-1):
        if runs[i].text==a and runs[i+1].text==b: return i
    return -1

def _para_contains(p: Paragraph, token: str) -> bool:
    return token in p.text

# --- rodapé da 1ª página (Estado - Brasil -> {estado}) ---
def _replace_footer_estado_first_page(doc: Document, estado: str):
    estado = (estado or "").strip()
    if not estado: return
    try: sec = doc.sections[0]
    except: return
    foot = getattr(sec,"first_page_footer",None) if getattr(sec,"different_first_page_header_footer",False) else sec.footer
    if not foot: return
    for p in foot.paragraphs:
        if "Brasil" not in p.text: continue
        # tentativa cirúrgica
        for r in p.runs:
            if r.text.strip()=="Estado":
                r.text = estado
                break
        else:
            # fallback: reescreve apenas o texto (mantém parágrafo)
            t_nodes = [n for n in p._p.iter() if str(getattr(n,"tag","")).endswith("}t")]
            if not t_nodes: continue
            original = "".join((n.text or "") for n in t_nodes)
            bound = original.find("Brasil")
            idx = original.rfind("Estado", 0, bound if bound!=-1 else len(original))
            if idx!=-1:
                new = original[:idx] + estado + original[idx+6:]
                t_nodes[0].text = new
                for tn in t_nodes[1:]: tn.text = ""

# --- remover highlight em todo doc (corpo + headers/footers) ---
def _remove_highlight(doc: Document):
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.highlight_color: r.font.highlight_color=None
    for sec in doc.sections:
        for area in (getattr(sec,"header",None), getattr(sec,"footer",None),
                     getattr(sec,"first_page_header",None), getattr(sec,"first_page_footer",None)):
            if area and hasattr(area,"paragraphs"):
                for p in area.paragraphs:
                    for r in p.runs:
                        if r.font.highlight_color: r.font.highlight_color=None

# --- preenchimento run-a-run da capa/preâmbulo ---
def fill_first_page(template_bytes: bytes, ctx: Dict[str,str]) -> bytes:
    doc = Document(BytesIO(template_bytes))
    pars = _iter_paragraphs_before_sumario(doc)
    V = lambda k: (ctx.get(k) or "").strip()

    for p in pars:
        runs = p.runs
        if not runs: continue

        # Título
        i = _first_run_eq(runs,"Nome da Empresa")
        if i>=0 and V("empresa"): runs[i].text = V("empresa")

        # Empresa: Aaa
        i = _first_run_eq(runs,"Empresa: ")
        if i>=0 and i+1<len(runs) and runs[i+1].text=="Aaa" and V("empresa"):
            runs[i+1].text = V("empresa")

        # CNPJ
        if V("cnpj"):
            j = _first_run_eq(runs,"XX.XXX.XXX/XXXX-XX")
            if j>=0: runs[j].text = V("cnpj")
            else:
                j = _find_pair(runs,"XX",".XXX.XXX/XXXX-XX")
                if j>=0: runs[j].text = V("cnpj"); runs[j+1].text = ""

        # XXXX/XXXX (regra: se janela ±8 runs contém “contestação” -> CT; senão -> TEC)
        def _apply_xxxx(i_num:int, j_slash:Optional[int]=None, k_num2:Optional[int]=None):
            if not (V("parecer") or V("parecer_tecnico")): return
            start = max(0, i_num-8); end = min(len(runs),(k_num2 if k_num2 is not None else i_num)+9)
            win = _norm_ascii_lower("".join(r.text for r in runs[start:end]))
            use = V("parecer") if "contestacao" in win and V("parecer") else V("parecer_tecnico")
            if not use: return
            if j_slash is None: runs[i_num].text = use
            else: runs[i_num].text = use; runs[j_slash].text = ""; runs[k_num2].text = ""
        j1 = _first_run_eq(runs,"XXXX/XXXX")
        while j1>=0: _apply_xxxx(j1); j1 = _first_run_eq(runs,"XXXX/XXXX",start=j1+1)
        i3=0
        while i3<len(runs)-2:
            if runs[i3].text=="XXXX" and runs[i3+1].text=="/" and runs[i3+2].text=="XXXX":
                _apply_xxxx(i3,i3+1,i3+2); i3+=3
            else: i3+=1

        # Ano-base 20XX (troca só o "XX")
        if V("ano_base_2"):
            low = p.text.lower()
            for k in range(1,len(runs)):
                if runs[k].text=="XX" and runs[k-1].text.endswith("20") and ("ano-base" in low):
                    runs[k].text = V("ano_base_2")

        # Parágrafo “A empresa …”
        if _para_contains(p,"A empresa"):
            if V("empresa"):
                j = _first_run_eq(runs,"Razão Social")
                if j>=0: runs[j].text = V("empresa")
            if V("cnpj"):
                j = _first_run_eq(runs,"XX.XXX.XXX/XXXX-XX")
                if j>=0: runs[j].text = V("cnpj")
            if V("logradouro"):
                j = _first_run_eq(runs,"Rua/Avenida/Estrada")
                if j>=0: runs[j].text = V("logradouro")
            if V("numero"):
                j = _first_run_eq(runs,"Número")
                if j>=0: runs[j].text = V("numero")
            if V("cep_5") or V("cep_3"):
                j5 = _first_run_eq(runs,"XXXX")
                if j5>=0 and j5+2<len(runs) and runs[j5+1].text=="-" and runs[j5+2].text=="XXX" and "CEP" in p.text:
                    if V("cep_5"): runs[j5].text = V("cep_5")
                    if V("cep_3"): runs[j5+2].text = V("cep_3")
            if V("municipio") or V("uf"):
                jc = _first_run_eq(runs,"Cidade"); js = _first_run_eq(runs," – "); ju = _first_run_eq(runs,"Estado")
                if jc>=0 and js>=0 and ju>=0:
                    if V("municipio"): runs[jc].text = V("municipio")
                    if V("uf"): runs[ju].text = V("uf")

        # Linha de data
        if ("Estado" in p.text) and ("Mês" in p.text):
            if V("estado_nome"):
                j = _first_run_eq(runs,"Estado")
                if j>=0: runs[j].text = V("estado_nome")
            if V("dia2"):
                for k in range(1,len(runs)):
                    if runs[k].text=="XX": runs[k].text = V("dia2"); break
            if V("mes_ext"):
                j = _first_run_eq(runs,"Mês")
                if j>=0: runs[j].text = V("mes_ext")
            if V("ano_data_2"):
                for k in range(1,len(runs)):
                    if runs[k].text=="XX" and runs[k-1].text.endswith("20"):
                        runs[k].text = V("ano_data_2")

        # Lote
        if V("lote"):
            i = _first_run_eq(runs,"Número do lote de liberação: ")
            if i>=0 and i+1<len(runs) and runs[i+1].text=="X": runs[i+1].text = V("lote")

        # Responsável
        if V("responsavel"):
            j = _first_run_eq(runs,"Nome Responsável (conforme cadastro no FormPD)")
            if j>=0: runs[j].text = V("responsavel")

    # Rodapé da 1ª página e highlight
    _replace_footer_estado_first_page(doc, V("estado_nome"))
    _remove_highlight(doc)

    out = BytesIO(); doc.save(out); out.seek(0); return out.read()